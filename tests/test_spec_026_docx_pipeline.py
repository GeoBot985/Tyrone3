import os
import tempfile
import unittest
from unittest.mock import MagicMock, patch

from docx import Document

from rag.docx_extractor import extract_docx_structured
from rag.ingest import ingest_document


class TestSpec026DocxPipeline(unittest.TestCase):
    def _temp_docx(self) -> str:
        handle = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        handle.close()
        self.addCleanup(lambda: os.path.exists(handle.name) and os.remove(handle.name))
        return handle.name

    def test_docx_extractor_reads_paragraphs_and_tables(self):
        path = self._temp_docx()
        doc = Document()
        doc.add_paragraph("Executive summary paragraph.")
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Member"
        table.rows[0].cells[1].text = "Amount"
        table.rows[1].cells[0].text = "George"
        table.rows[1].cells[1].text = "100"
        doc.save(path)

        result = extract_docx_structured(path)

        self.assertTrue(result["success"])
        self.assertEqual(result["method"], "python_docx_structured")
        self.assertEqual(result["paragraph_count"], 1)
        self.assertEqual(result["table_count"], 1)
        self.assertEqual(result["table_row_count"], 1)
        self.assertIn("Executive summary paragraph.", result["text"])
        self.assertIn("[DOCX | Block: 2 | Ref: Table 1 Row 2 | Region: table_row]", result["text"])

    @patch("rag.ingest.insert_chunk")
    @patch("rag.ingest.insert_document")
    @patch("rag.ingest.embed_text", return_value=[0.1, 0.2, 0.3])
    def test_ingest_docx_uses_structured_pipeline(self, _mock_embed, _mock_insert_document, _mock_insert_chunk):
        path = self._temp_docx()
        doc = Document()
        doc.add_paragraph("Policy update paragraph.")
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Owner"
        table.rows[0].cells[1].text = "Status"
        table.rows[1].cells[0].text = "Nina"
        table.rows[1].cells[1].text = "Approved"
        doc.save(path)

        result = ingest_document(path, MagicMock(), document_name="policy.docx")

        self.assertEqual(result["file_type"], "docx")
        self.assertEqual(result["ingestion_method"], "python_docx_structured")
        self.assertEqual(result["provenance"]["primary_extractor"], "python_docx_structured")
        self.assertTrue(result["provenance"]["has_table_rows"])
        self.assertEqual(result["provenance"]["region_counts"].get("table_row"), 1)
        self.assertGreater(result["chunk_count"], 0)

    def test_docx_failure_is_clear(self):
        handle = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        handle.write(b"not-a-real-docx")
        handle.close()
        self.addCleanup(lambda: os.path.exists(handle.name) and os.remove(handle.name))

        with self.assertRaisesRegex(Exception, "Failed to read DOCX file") as ctx:
            ingest_document(handle.name, MagicMock(), document_name="broken.docx")

        self.assertEqual(ctx.exception.provenance["failure_stage"], "extract_docx")
        self.assertEqual(ctx.exception.provenance["primary_extractor"], "python_docx_structured")
