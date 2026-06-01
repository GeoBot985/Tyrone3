from __future__ import annotations

import csv
import json
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any

import fitz
from docx import Document
from openpyxl import Workbook
from PIL import Image, ImageDraw, ImageFont

from eval.common import (
    CORPUS_DIR,
    GOLDEN_PATH,
    ensure_eval_dirs,
    load_jsonl,
    stable_doc_id,
    write_jsonl,
)


@dataclass(frozen=True)
class CorpusDoc:
    file_name: str
    kind: str
    title: str
    content: Any


CORPUS_DOCS = [
    CorpusDoc(
        file_name="alpha_policy.txt",
        kind="txt",
        title="Alpha office policy",
        content=(
            "Alpha office policy\n"
            "Lunch stipend is capped at 120 ZAR per day.\n"
            "Friday standup starts at 09:30 in Room C.\n"
            "The security desk extension is 404.\n"
            "The door code changes every Monday at 07:00.\n"
            "Do not leave laptop chargers overnight.\n"
            "Desk booking uses TeamUp.\n"
        ),
    ),
    CorpusDoc(
        file_name="beta_meeting.docx",
        kind="docx",
        title="Beta sprint meeting notes",
        content={
            "paragraphs": [
                "Beta sprint meeting notes for the portfolio release.",
                "Decision: the corpus will include PDF, DOCX, XLSX, CSV, and TXT files.",
                "Budget approved: 18 hours.",
                "Demo date: Tuesday 14 June 2026.",
            ],
            "table": [
                ["Owner", "Task", "Due date"],
                ["Nomvula", "Validate OCR on the scanned PDF", "Thursday 2 June 2026"],
                ["Musa", "Update the README", "Friday 3 June 2026"],
                ["Aisha", "Prepare screenshots", "Saturday 4 June 2026"],
            ],
        },
    ),
    CorpusDoc(
        file_name="gamma_metrics.csv",
        kind="csv",
        title="Gamma sales metrics",
        content=[
            ["product", "trials", "paid_users", "conversion_pct", "revenue_zar"],
            ["Alpha", "120", "18", "15.0", "36000"],
            ["Beta", "90", "27", "30.0", "54000"],
            ["Gamma", "150", "30", "20.0", "60000"],
            ["Delta", "70", "7", "10.0", "14000"],
        ],
    ),
    CorpusDoc(
        file_name="delta_schedule.xlsx",
        kind="xlsx",
        title="Delta schedule workbook",
        content={
            "Roster": [
                ["Name", "Role", "Shift"],
                ["Nandi", "QA", "Morning"],
                ["Joel", "DevOps", "Afternoon"],
                ["Priya", "Finance", "Evening"],
            ],
            "Milestones": [
                ["Milestone", "Owner", "Due Date", "Status"],
                ["Alpha Freeze", "Theo", date(2026, 6, 1), "done"],
                ["Beta Review", "Lindiwe", date(2026, 6, 3), "open"],
                ["Gamma Release", "Sam", date(2026, 6, 7), "queued"],
            ],
        },
    ),
    CorpusDoc(
        file_name="epsilon_scanned.pdf",
        kind="scanned_pdf",
        title="Epsilon scanned field note",
        content=[
            "Epsilon field note for the archive scanner and OCR verification.",
            "Emergency kit code is Indigo-7.",
            "Backup key lives in drawer B4.",
            "Server room escort number is 771.",
            "Cabin log closes at 18:00.",
            "Visitor log requires a signature before entry.",
            "North corridor access is closed after 18:00.",
            "Please keep this page aligned for reliable OCR capture.",
            "The archive shelf label is E-14.",
            "Record the inspection time as 17:45.",
            "If the seal is broken, notify the duty lead immediately.",
            "The backup battery shelf is marked Delta Two.",
            "Use the red marker for the final sign-off line.",
        ],
    ),
    CorpusDoc(
        file_name="zeta_brief.pdf",
        kind="pdf",
        title="Zeta travel brief",
        content=(
            "Zeta travel brief for June client work and finance review.\n"
            "Travel allowance is 1800 ZAR.\n"
            "Hotel ceiling is 950 ZAR.\n"
            "Finance approvals go to Priya.\n"
            "Client visit date is 12 June 2026.\n"
            "Meal reimbursement requires receipts.\n"
            "Carry a printed itinerary for the site visit.\n"
            "Keep the taxi receipt stub in the folder.\n"
            "Submit the expense form within three working days.\n"
            "The rail backup option is approved if flights are delayed.\n"
            "If the trip changes, notify finance before noon.\n"
            "Use the blue folder for all signed paperwork.\n"
            "The office code for reimbursement follow-up is 27B.\n"
        ),
    ),
]


def _write_txt(path: Path, text: str) -> None:
    path.write_text(text, encoding="utf-8")


def _write_csv(path: Path, rows: list[list[str]]) -> None:
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle)
        writer.writerows(rows)


def _write_docx(path: Path, spec: dict[str, Any]) -> None:
    document = Document()
    document.add_heading("Beta sprint meeting notes", level=1)
    for paragraph in spec["paragraphs"]:
        document.add_paragraph(paragraph)
    table_rows = spec["table"]
    table = document.add_table(rows=1, cols=len(table_rows[0]))
    for cell, header in zip(table.rows[0].cells, table_rows[0]):
        cell.text = header
    for row in table_rows[1:]:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = str(value)
    document.save(path)


def _write_xlsx(path: Path, sheets: dict[str, list[list[Any]]]) -> None:
    workbook = Workbook()
    first = True
    for sheet_name, rows in sheets.items():
        if first:
            worksheet = workbook.active
            worksheet.title = sheet_name
            first = False
        else:
            worksheet = workbook.create_sheet(title=sheet_name)
        for row in rows:
            worksheet.append(row)
    workbook.save(path)


def _write_pdf(path: Path, text: str) -> None:
    image = Image.new("RGB", (1800, 2400), "white")
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.truetype("arial.ttf", 42)
    except OSError:
        font = ImageFont.load_default()
    y = 100
    for line in text.splitlines():
        draw.text((100, y), line, fill="black", font=font)
        y += 110
    image.save(path, "PDF", resolution=200.0)


def _write_scanned_pdf(path: Path, lines: list[str]) -> None:
    image = Image.new("RGB", (1600, 2000), "white")
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.truetype("arial.ttf", 44)
    except OSError:
        font = ImageFont.load_default()
    y = 100
    for line in lines:
        draw.text((100, y), line, fill="black", font=font)
        y += 120
    image.save(path, "PDF", resolution=200.0)


def build_corpus(overwrite: bool = True) -> list[dict[str, str]]:
    ensure_eval_dirs()
    doc_index: list[dict[str, str]] = []

    for spec in CORPUS_DOCS:
        path = CORPUS_DIR / spec.file_name
        if path.exists() and not overwrite:
            doc_index.append({"file_name": spec.file_name, "doc_id": stable_doc_id(path)})
            continue

        if spec.kind == "txt":
            _write_txt(path, str(spec.content))
        elif spec.kind == "docx":
            _write_docx(path, spec.content)
        elif spec.kind == "csv":
            _write_csv(path, spec.content)
        elif spec.kind == "xlsx":
            _write_xlsx(path, spec.content)
        elif spec.kind == "pdf":
            _write_pdf(path, str(spec.content))
        elif spec.kind == "scanned_pdf":
            _write_scanned_pdf(path, spec.content)
        else:
            raise ValueError(f"Unsupported corpus kind: {spec.kind}")

        doc_index.append({"file_name": spec.file_name, "doc_id": stable_doc_id(path)})

    return doc_index


def _cases_for_docs(doc_index: list[dict[str, str]]) -> list[dict[str, Any]]:
    doc_ids = {item["file_name"]: item["doc_id"] for item in doc_index}
    cases: list[dict[str, Any]] = []

    def add_case(
        question: str,
        file_names: list[str],
        expected_substrings: list[str],
        must_contain: list[str],
        must_not_contain: list[str] | None = None,
        should_refuse: bool = False,
        mode: str = "document",
        faithfulness_eval: bool = False,
        confidence_eval: bool = False,
    ) -> None:
        cases.append(
            {
                "question": question,
                "mode": mode,
                "expected_doc_ids": [doc_ids[name] for name in file_names],
                "expected_chunk_substrings": expected_substrings,
                "answer_must_contain": must_contain,
                "answer_must_not_contain": must_not_contain or ["Insufficient information"],
                "should_refuse": should_refuse,
                "faithfulness_eval": faithfulness_eval,
                "confidence_eval": confidence_eval or faithfulness_eval or should_refuse,
            }
        )

    policy = "alpha_policy.txt"
    add_case(
        "What is the lunch stipend cap?",
        [policy],
        ["120 ZAR per day"],
        ["120", "ZAR"],
        faithfulness_eval=True,
        confidence_eval=True,
    )
    add_case("When does the Friday standup start?", [policy], ["Friday standup starts at 09:30"], ["09:30"])
    add_case("What extension reaches the security desk?", [policy], ["security desk extension is 404"], ["404"])
    add_case(
        "Which system is used for desk booking?",
        [policy],
        ["Desk booking uses TeamUp"],
        ["TeamUp"],
        faithfulness_eval=True,
        confidence_eval=True,
    )
    add_case("What day does the door code change?", [policy], ["door code changes every Monday at 07:00"], ["Monday", "07:00"])
    add_case("What should not be left overnight?", [policy], ["Do not leave laptop chargers overnight"], ["laptop chargers"])
    add_case(
        "What is the policy on laptop chargers and desk booking?",
        [policy],
        ["Do not leave laptop chargers overnight", "Desk booking uses TeamUp"],
        ["chargers", "TeamUp"],
    )

    meeting = "beta_meeting.docx"
    add_case(
        "Who prepares screenshots?",
        [meeting],
        ["Aisha", "Prepare screenshots"],
        ["Aisha"],
        faithfulness_eval=True,
        confidence_eval=True,
    )
    add_case(
        "What is the demo date?",
        [meeting],
        ["Demo date: Tuesday 14 June 2026"],
        ["14 June 2026"],
        faithfulness_eval=True,
        confidence_eval=True,
    )
    add_case("How many hours were budgeted?", [meeting], ["Budget approved: 18 hours"], ["18 hours"])
    add_case(
        "Which file types are included in the corpus decision?",
        [meeting],
        ["PDF, DOCX, XLSX, CSV, and TXT"],
        ["PDF", "DOCX", "XLSX", "CSV", "TXT"],
    )
    add_case("Who validates OCR on the scanned PDF?", [meeting], ["Nomvula", "Validate OCR"], ["Nomvula"])
    add_case("What does Musa update?", [meeting], ["Musa", "Update the README"], ["README"])
    add_case(
        "Which task is due on Saturday 4 June 2026?",
        [meeting],
        ["Saturday 4 June 2026"],
        ["Aisha", "screenshots"],
    )

    metrics = "gamma_metrics.csv"
    add_case(
        "Which product had the highest conversion?",
        [metrics],
        ["Beta", "30.0"],
        ["Beta", "30.0"],
        faithfulness_eval=True,
        confidence_eval=True,
    )
    add_case(
        "What revenue did Gamma generate?",
        [metrics],
        ["Gamma", "60000"],
        ["60000"],
        faithfulness_eval=True,
        confidence_eval=True,
    )
    add_case("How many paid users did Alpha have?", [metrics], ["Alpha", "18"], ["18"])
    add_case("Which product had 7 paid users?", [metrics], ["Delta", "7"], ["Delta"])
    add_case("Which product had 120 trials?", [metrics], ["Alpha", "120"], ["Alpha"])
    add_case("What was Beta's conversion percentage?", [metrics], ["Beta", "30.0"], ["30.0"])
    add_case("What revenue did Delta generate?", [metrics], ["Delta", "14000"], ["14000"])

    schedule = "delta_schedule.xlsx"
    add_case(
        "Who owns Gamma Release?",
        [schedule],
        ["Gamma Release", "Sam"],
        ["Sam"],
        faithfulness_eval=True,
        confidence_eval=True,
    )
    add_case(
        "What is the due date of Beta Review?",
        [schedule],
        ["Beta Review", "03 June 2026"],
        ["3 June 2026", "03 June 2026"],
        faithfulness_eval=True,
        confidence_eval=True,
    )
    add_case("Which role does Nandi have?", [schedule], ["Nandi", "QA"], ["QA"])
    add_case("What shift is Joel on?", [schedule], ["Joel", "Afternoon"], ["Afternoon"])
    add_case("Who owns Alpha Freeze?", [schedule], ["Alpha Freeze", "Theo"], ["Theo"])
    add_case("What status does Gamma Release have?", [schedule], ["Gamma Release", "queued"], ["queued"])
    add_case("What due date is listed for Alpha Freeze?", [schedule], ["Alpha Freeze", "01 June 2026"], ["1 June 2026", "01 June 2026"])

    scanned = "epsilon_scanned.pdf"
    add_case(
        "What is the backup key drawer?",
        [scanned],
        ["drawer B4"],
        ["B4"],
        faithfulness_eval=True,
        confidence_eval=True,
    )
    add_case(
        "What is the emergency kit code?",
        [scanned],
        ["Indigo-7"],
        ["Indigo-7"],
        faithfulness_eval=True,
        confidence_eval=True,
    )
    add_case("What is the server room escort number?", [scanned], ["771"], ["771"])
    add_case("When does the cabin log close?", [scanned], ["Cabin log closes at 18:00"], ["18:00"])
    add_case("Where does the backup key live?", [scanned], ["Backup key lives in drawer B4"], ["drawer B4"])
    add_case("What code is printed on the field note?", [scanned], ["Indigo-7"], ["Indigo-7"])

    brief = "zeta_brief.pdf"
    add_case(
        "What is the travel allowance?",
        [brief],
        ["Travel allowance is 1800 ZAR"],
        ["1800", "ZAR"],
        faithfulness_eval=True,
        confidence_eval=True,
    )
    add_case(
        "What is the hotel ceiling?",
        [brief],
        ["Hotel ceiling is 950 ZAR"],
        ["950", "ZAR"],
    )
    add_case(
        "Who receives finance approvals?",
        [brief],
        ["Finance approvals go to Priya"],
        ["Priya"],
        faithfulness_eval=True,
        confidence_eval=True,
    )
    add_case("What is the client visit date?", [brief], ["Client visit date is 12 June 2026"], ["12 June 2026"])
    add_case("What reimbursal rule is listed for meals?", [brief], ["Meal reimbursement requires receipts"], ["receipts"])
    add_case("What is the travel allowance and hotel ceiling?", [brief], ["1800 ZAR", "950 ZAR"], ["1800", "950"])
    add_case("Who should approve the travel brief?", [brief], ["Priya"], ["Priya"])

    add_case(
        "Across the office policy and the travel brief, what are the two ZAR caps?",
        [policy, brief],
        ["120 ZAR per day", "Travel allowance is 1800 ZAR", "Hotel ceiling is 950 ZAR"],
        ["120 ZAR", "1800 ZAR", "950 ZAR"],
        faithfulness_eval=True,
        confidence_eval=True,
    )
    add_case(
        "Between the meeting notes and the schedule workbook, who owns screenshots and Gamma Release?",
        [meeting, schedule],
        ["Aisha", "Gamma Release", "Sam"],
        ["Aisha", "Sam"],
        faithfulness_eval=True,
        confidence_eval=True,
    )
    add_case(
        "Which document names the desk booking system and who receives finance approvals?",
        [policy, brief],
        ["Desk booking uses TeamUp", "Finance approvals go to Priya"],
        ["TeamUp", "Priya"],
    )
    add_case(
        "Which docs mention June 2026 dates for the demo and client visit?",
        [meeting, brief],
        ["Demo date: Tuesday 14 June 2026", "Client visit date is 12 June 2026"],
        ["14 June 2026", "12 June 2026"],
    )

    refuse_questions = [
        "What is the warranty period for the lunar orchid harvester?",
        "Who signed the Zephyr cinema contract last quarter?",
        "What are the emission standards for the moon aquarium reactor?",
        "How many gondolas are required for the sapphire telescope launch?",
        "Which ice cream tax rule applies to the polar submarine ordinance?",
        "What is the bird migration policy for the emerald satellite farm?",
        "Who manages the cinnamon drone library in Oslo?",
        "What is the approval code for the velvet cactus treaty?",
    ]
    for question in refuse_questions:
        cases.append(
            {
                "question": question,
                "mode": "document",
                "expected_doc_ids": [],
                "expected_chunk_substrings": [],
                "answer_must_contain": ["Insufficient information"],
                "answer_must_not_contain": ["120 ZAR", "TeamUp", "Indigo-7"],
                "should_refuse": True,
            }
        )

    return cases


def build_golden(doc_index: list[dict[str, str]]) -> list[dict[str, Any]]:
    cases = _cases_for_docs(doc_index)
    write_jsonl(GOLDEN_PATH, cases)
    return cases


def main() -> int:
    doc_index = build_corpus(overwrite=True)
    build_golden(doc_index)
    manifest_path = CORPUS_DIR / "manifest.json"
    manifest_path.write_text(json.dumps(doc_index, indent=2, sort_keys=True), encoding="utf-8")
    print(f"Wrote {len(doc_index)} corpus documents to {CORPUS_DIR}")
    print(f"Wrote {len(load_jsonl(GOLDEN_PATH))} golden cases to {GOLDEN_PATH}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
